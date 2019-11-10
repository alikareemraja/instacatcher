#!/bin/python
"""
Hello World, but with more meat.
"""


import instaloader
import os
import wx
from datetime import datetime
import sys
from docx import Document
from docx.shared import Inches
import xlsxwriter
import xlrd
import calendar
import regex
import emoji
from pathlib import Path
from threading import *

class InstaLoaderThread(Thread):
    
    """Worker Thread Class."""
    def __init__(self, progress_bar, state, influencer):
        """Init Worker Thread Class."""
        Thread.__init__(self)
        self.influencer = influencer;
        self.progress_bar = progress_bar
        self._want_abort = 0
        # This starts the thread running on creation, but you could
        # also make the GUI thread responsible for calling this
        self.state = state;
        self.start()
    

    def run(self):

        username = self.influencer
        
        L = instaloader.Instaloader()

        self.post_folder = username + '_posts';
        self.stories_folder = username + '_stories';

        auth = L.login(self.state.login_user,self.state.login_password)
        self.count_posts = 0

        try: # try to load user profile; else give error console text
            self.profile = instaloader.Profile.from_username(L.context, username)  #self.state.usrOfPosts
        except:
            #self._application_window.SetStatusText('User does not exist. Try different User name.')
            return

        if self.state.isDate == False:
            #self._application_window.SetStatusText('Date format is not recognized. Please change to correct format.')
            return

        excelData = []

        if self.state.timeFrom > datetime.today().date(): # lower bound needs to be earlier than today
            self.state.timeFrom = datetime.strptime("1900-01-01", '%Y-%m-%d').date()

        if self.state.timeFrom > self.state.timeTo: # lower bound needs to be earlier than upper bound
            self.state.timeFrom = datetime.strptime("1900-01-01", '%Y-%m-%d').date()
            if self.state.timeTo <= datetime.strptime("1900-01-01", '%Y-%m-%d').date():
                self.state.timeTo = datetime.today().date()

        storiesExcelData = []

        if self.state.getStories:
            stories = L.get_stories([self.profile.userid]) 
            
            for story in stories:
                    for item in story.get_items():
                        storiesExcelData.append([item.date_local, calendar.day_name[item.date_local.weekday()], item.owner_profile.followers, "Video" if item.is_video == True else "Photo"])                    
                        print("Downloading Story item" )
                        L.download_storyitem(item, self.stories_folder)
                        #self.count_posts = self.count_posts + 1
                        self.printProgress()

                
        """Load Posts:"""
        #self._application_window.SetStatusText('Loading posts...')
        posts = self.profile.get_posts()

        # counter variable for looping through posts
        i = 0
        
        
        if not self.state.savePosts:
            wx.CallAfter(self.progress_bar.SetLabel, "Done - "+ str(self.count_posts) +" Items")
            return

        for post in posts: # iterate through each post
            

            if self._want_abort:
                # Use a result of None to acknowledge the abort (of
                # course you can use whatever you'd like or even
                # a separate event type)
                return
            
            if post.date_utc.date() >= self.state.timeFrom and post.date_utc.date() <= self.state.timeTo:
                i += 1
            else:
                #post = (next(posts))
                #if post.date_utc.date() < self.state.timeFrom:
                #    i += 10000
                #    post = (next(posts))
                continue

            trueFalse = False
            try:
                trueFalse = L.download_post(post, self.post_folder)
            except:
                trueFalse = False

            #if trueFalse:
                #self._application_window.SetStatusText('Downloaded post: ' + post.caption)
            #else:
                #self._application_window.SetStatusText('/!\ Error downloading posts.')

            if not self.state.createDocs:
                self.printProgress()
                continue

            postTime = post.date_utc
            fileName = self.getDate(postTime)

            project_path = os.path.dirname(sys.modules['__main__'].__file__)
            #script_dir = os.path.dirname(__file__)  # <-- absolute dir the script is in
            rel_path = os.path.join(self.post_folder , "{0}.txt".format(fileName))
            abs_file_path = os.path.join(project_path, rel_path)

            try:
                caption = open(abs_file_path, 'r+').read()
            except:
                print("Image has no caption.")


            """Catch information for Excel Sheet:"""
            postWeekday = calendar.day_name[postTime.weekday()]

            followers = 0
            try:
                followers = self.profile.followers
            except:
                followers = 1
                print("Error catching FOLLOWERS of Influencer. FOLLOWERS set to 1.")

            capLength = 0
            capEmoticons = 0
            cap = post.caption
            
            if cap != None:
                capLength = len(cap)
                data = regex.findall(r'\X', cap)
                for word in data:
                    if any(char in emoji.UNICODE_EMOJI for char in word):
                        capEmoticons += 1
            else:
                capLength = 0
                capEmoticons = 0

            capHashtags = 0
            try:
                hashtags = post.caption_hashtags
                capHashtags = len(hashtags)
            except:
                capHashtags = 0
                print("Hashtags could not be fetched")


            likes = 0
            try:
                likes = post.likes
            except:
                likes = 0
                print("Error catching LIKES of Influencer.")

            likesPerFollower = likes / followers



            """Erstelle .docx Dokument:"""
            document = Document()

            # Überschrift erstellen:
            mediaType=""
            if post.is_video:
                mediaType="Video (Dauer: %s Sek.)" % str(post.video_duration)
            else:
                mediaType="Foto"

            document.add_heading('{0} von {1} am {2}:'.format(mediaType,self.post_folder,fileName), level=1)
            document.add_paragraph('Information caught at %s' % datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
            
            if os.path.exists(os.path.join(self.post_folder , "{0}.jpg".format(fileName))):
                document.add_picture('{0}/{1}.jpg'.format(self.post_folder,fileName), width=Inches(2.5))
            else:
                counter = 1
                while os.path.exists(os.path.join(self.post_folder , "{0}_{1}.jpg".format(fileName,counter))):
                    document.add_picture(os.path.join(self.post_folder , "{0}_{1}.jpg".format(fileName,counter)), width=Inches(2.5))
                    counter += 1

            mediaTypeExcel = post.typename
            if mediaTypeExcel == "GraphSidecar":
                sidecars = post.get_sidecar_nodes()
                counter = 0
                for _ in sidecars:
                    counter += 1
                mediaTypeExcel = "Fotoserie (%s)" % str(counter)
            else:
                mediaTypeExcel = mediaType

            # Likes:
            document.add_heading('%s Likes' % str(post.likes), level=3)

            # Bildunterschrift:
            document.add_heading('Caption:', level=3)
            document.add_paragraph(caption)

            # Iterate over all comments of current post
            document.add_heading('%s Comments:' % str(post.comments), level=3)
            r = document.add_paragraph()
            comments = post.get_comments()


            likesPerFollower = likes / followers

            com0To1 = 0
            com1To3 = 0
            com3To12 = 0
            com12To24 = 0
            com24To48 = 0
            comRest = 0
            comEmoticonCounter = 0

            try:
                for com in comments:
                    q = document.add_paragraph()
                    q.add_run("%s: " % str(getattr(com, 'created_at_utc'))).bold
                    q.add_run(com.text)

                    # check time difference of each comment to post
                    comTime = getattr(com, 'created_at_utc')
                    timeDifference = (comTime - postTime).total_seconds()

                    if timeDifference <= 3600:
                        com0To1 += 1
                    elif timeDifference <= 10800:
                        com1To3 +=1
                    elif timeDifference <= 43200:
                        com3To12 += 1
                    elif timeDifference <= 86400:
                        com12To24 += 1
                    elif timeDifference <= 172800:
                        com24To48 += 1
                    else:
                        comRest += 1

                    data = regex.findall(r'\X', getattr(com, 'text'))
                    for word in data:
                        if any(char in emoji.UNICODE_EMOJI for char in word):
                            comEmoticonCounter += 1

                    # capHashtags = 0
            except:
                print("SOMETHING WENT WRONG!")
                pass
            

            comCount = post.comments
            comAnswers = comCount - (com0To1 + com1To3 + com3To12 + com12To24 + com24To48 + comRest)
            comPerFollower = comCount / followers

            if comCount != 0:
                comEmoticonAverage = comEmoticonCounter / comCount
            else:
                comEmoticonAverage = 0

            """Create EXCEL Data and store to Array"""
            dataPoint = [postTime, postWeekday, followers, mediaTypeExcel, capLength, capEmoticons, capHashtags, likes, likesPerFollower, comCount, comPerFollower, com0To1, com1To3, com3To12, com12To24, com24To48, comRest, comAnswers, comEmoticonAverage]
            excelData.append(dataPoint)

            """Create folder if not existent:"""
            try:
                os.mkdir("data_%s" % username)
                print("Directory data_", username, " created.")
            except:
                print("Directory data_", username, " already exists.")

            document.save(os.path.join("data_{0}".format(username) , "{0}.docx".format(fileName)))

            self.printProgress()
            
            #try:
            #    post = (next(posts))
            #except:
            #    print("No more posts to iterate.")
            #    #self._application_window.SetStatusText('Download Finished')
            #    break;


        wx.CallAfter(self.progress_bar.SetLabel, "Writing Summary")

        """Create and built Excel Sheet:"""

        workbook = None;
        worksheets = [];
        rows = [2, 2]
        worksheet = None;
        sheets = []

        if os.path.exists(os.path.join("data_{0}".format(username) , "{0}.xlsx".format(username))):
            wbRD = xlrd.open_workbook(os.path.join("data_{0}".format(username) , "{0}.xlsx".format(username)))
            sheets = wbRD.sheets()
        
            
        workbook = xlsxwriter.Workbook(os.path.join("data_{0}".format(username) , "{0}.xlsx".format(username)), {'constant_memory': True})
        worksheets.append(workbook.add_worksheet())

        bold = workbook.add_format({'bold': 1})
        bold = workbook.add_format({'bold': 2})
        date_format = workbook.add_format({'num_format': 'yyyy/mm/dd hh:mm'})
        merge_format = workbook.add_format({'align': 'center'})

        #create header
        worksheets[0].write('A1', 'PostTime', bold)
        worksheets[0].write('B1', 'Weekday', bold)
        worksheets[0].write('C1', 'Followers', bold)
        worksheets[0].write('D1', 'MediaType', bold)

        worksheets[0].merge_range('E1:G1', 'Caption', merge_format)
        worksheets[0].merge_range('H1:I1', 'Likes', merge_format)
        worksheets[0].merge_range('J1:S1', 'Comments', merge_format)
        worksheets[0].write('E2', 'Length', bold)
        worksheets[0].write('F2', 'Emojis', bold)
        worksheets[0].write('G2', 'Hashtags', bold)
        worksheets[0].write('H2', 'Count', bold)
        worksheets[0].write('I2', 'per Follower', bold)
        worksheets[0].write('J2', 'Count', bold)
        worksheets[0].write('K2', 'per Follower', bold)
        worksheets[0].write('L2', '0-1 h', bold)
        worksheets[0].write('M2', '1-3 h', bold)
        worksheets[0].write('N2', '3-12 h', bold)
        worksheets[0].write('O2', '12-24 h', bold)
        worksheets[0].write('P2', '24-48 h', bold)
        worksheets[0].write('Q2', 'Rest', bold)
        worksheets[0].write('R2', 'Comment Answers', bold)
        worksheets[0].write('S2', 'Emojis p. Comment', bold)

        worksheets.append(workbook.add_worksheet())
        bold = workbook.add_format({'bold': 1})
        bold = workbook.add_format({'bold': 2})
        date_format = workbook.add_format({'num_format': 'yyyy/mm/dd hh:mm'})
        merge_format = workbook.add_format({'align': 'center'})
        
        #create header
        worksheets[1].write('A1', 'PostTime', bold)
        worksheets[1].write('B1', 'Weekday', bold)
        worksheets[1].write('C1', 'Followers', bold)
        worksheets[1].write('D1', 'MediaType', bold)

        row = 2
        col = 0

        for element in excelData:
            
            worksheets[0].write_string(rows[0], col, str(element[0]))   # postTime
            worksheets[0].write_string(rows[0], col + 1, element[1])    # postWeekday
            worksheets[0].write_number(rows[0], col + 2, element[2])    # followers
            worksheets[0].write_string(rows[0], col + 3, element[3])    # mediaType
            worksheets[0].write_number(rows[0], col + 4, element[4])    # capLength
            worksheets[0].write_number(rows[0], col + 5, element[5])    # capEmoticons
            worksheets[0].write_number(rows[0], col + 6, element[6])    # capHashtags
            worksheets[0].write_number(rows[0], col + 7, element[7])    # likes
            worksheets[0].write_number(rows[0], col + 8, element[8])    # likesPerFollower
            worksheets[0].write_number(rows[0], col + 9, element[9])    # comCount
            worksheets[0].write_number(rows[0], col + 10, element[10])  # comPerFollower
            worksheets[0].write_number(rows[0], col + 11, element[11])  # com0To1
            worksheets[0].write_number(rows[0], col + 12, element[12])  # com1To3
            worksheets[0].write_number(rows[0], col + 13, element[13])  # com3To12
            worksheets[0].write_number(rows[0], col + 14, element[14])  # com12To24
            worksheets[0].write_number(rows[0], col + 15, element[15])  # com24To48
            worksheets[0].write_number(rows[0], col + 16, element[16])  # comRest
            worksheets[0].write_number(rows[0], col + 17, element[17])  # comAnswers
            worksheets[0].write_number(rows[0], col + 18, element[18])  # comEmoticonAverage

            rows[0] += 1

        row = 2
        col = 0

        for element in storiesExcelData:
            
            worksheets[1].write_string(rows[1], col, str(element[0]))   # postTime
            worksheets[1].write_string(rows[1], col + 1, element[1])    # postWeekday
            worksheets[1].write_number(rows[1], col + 2, element[2])    # followers
            worksheets[1].write_string(rows[1], col + 3, element[3])    # mediaType

            rows[1] += 1

        index = 0
        
        # run through the sheets and store sheets in workbook
        # this still doesn't write to the file yet
        if len(sheets) > 0:
            for sheet in sheets: # write data from old file
                #worksheets.append(workbook.add_worksheet(sheet.name))
                for row in (range(sheet.nrows - 2)):
                    for col in range(sheet.ncols):
                        worksheets[index].write(rows[index], col, sheet.cell(row + 2, col).value)
                    rows[index] = rows[index] + 1
                index = index + 1


        workbook.close()

        wx.CallAfter(self.progress_bar.SetLabel, "Done - "+ str(self.count_posts) +" Items")

        # Set Status Text if finished sucessfully
        print("Excel written for " + username)    
        #self._application_window.SetStatusText('Finished sucessfully.')                            

    def getDate (self, date: datetime):

        Y=""
        if date.year < 10:
           Y="0%s" % str(date.year)
        else:
            Y=str(date.year)

        M = ""
        if date.month < 10:
            M = "0%s" % str(date.month)
        else:
            M = str(date.month)

        D = ""
        if date.day < 10:
            D = "0%s" % str(date.day)
        else:
            D = str(date.day)

        h = ""
        if date.hour < 10:
            h = "0%s" % str(date.hour)
        else:
            h = str(date.hour)

        m = ""
        if date.minute < 10:
            m = "0%s" % str(date.minute)
        else:
            m = str(date.minute)

        s = ""
        if date.second < 10:
            s = "0%s" % str(date.second)
        else:
            s = str(date.second)

        return "%(Y)s-%(M)s-%(D)s_%(h)s-%(m)s-%(s)s_UTC" % {'Y': Y, 'M': M, 'D': D, 'h': h, 'm': m, 's': s}

    def abort(self):
        """abort worker thread."""
        # Method for use by main thread to signal an abort
        self._want_abort = 1

    def printProgress(self):
        self.count_posts = self.count_posts + 1
        wx.CallAfter(self.progress_bar.SetLabel, "In Progress - "+ str(self.count_posts) +" Items")